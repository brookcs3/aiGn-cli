"""
DOCX text extraction using python-docx
"""
import sys
from pathlib import Path

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text as a string
    """
    try:
        from docx import Document
    except ImportError:
        return "[ERROR] python-docx not installed. Run: pip install python-docx"

    path = Path(file_path)
    if not path.exists():
        return f"[ERROR] File not found: {file_path}"

    if path.suffix.lower() not in [".docx", ".doc"]:
        return f"[ERROR] Not a Word document: {file_path}"

    try:
        doc = Document(str(path))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    except Exception as e:
        return f"[ERROR] Failed to parse DOCX: {e}"


if __name__ == "__main__":
    # CLI usage for testing
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <file.docx>")
        sys.exit(1)

    text = extract_text_from_docx(sys.argv[1])
    print(text)
